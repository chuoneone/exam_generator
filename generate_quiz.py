import os
import re
import sys
import subprocess
import shutil
import json

# Make sure we can import docx
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("docx module is missing. Installing python-docx...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

# Define directories and file paths
TARGET_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "視力與偏見_20題_練習卷")
if not os.path.exists(TARGET_DIR):
    os.makedirs(TARGET_DIR)

student_html_path = os.path.join(TARGET_DIR, "視力與偏見_20題_學生點讀.html")
teacher_html_path = os.path.join(TARGET_DIR, "視力與偏見_20題_教師簡報.html")

# Template Paths
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "skills", "paper_exam_generator", "resources", "templates")
student_template_file = os.path.join(TEMPLATE_DIR, "student_template.html")
teacher_template_file = os.path.join(TEMPLATE_DIR, "teacher_template.html")

# Data Definition
article_title = "視力與偏見"
article_text = """在從紐約到波士頓的火車上，我發現我隔壁座的老先生是位盲人。

我的博士論文指導教授是位盲人，因此我和盲人談起話來，一點困難也沒有，我還弄了一杯熱騰騰的咖啡給他喝。

當時正值洛杉磯種族暴動的時期，我們的談話因此就談到了種族偏見的問題。

老先生告訴我，他是美國南方人，從小就認為黑人低人一等，他家的佣人是黑人，他在南方時從未和黑人一起吃過飯，也從未和黑人上過學，到了北方唸書，他有次被班上同學指定辦一次野餐會，他居然在請帖上註明「我們保留拒絕任何人的權利」。在南方這句話就是「我們不歡迎黑人」的意思，當時舉班譁然，他還被系主任抓去罵了一頓。

他說有時碰到黑人店員，付錢的時候，總將錢放在櫃台上，讓黑人拿去，而不肯和他的手有任何接觸。 

我笑著問他：「那你當然不會和黑人結婚了！」

he大笑起來：「我不和他們來往，如何會和黑人結婚？說實話，我當時認為任何白人和黑人結婚都會使父母蒙辱。」

可是，他在波士頓唸研究所的時候，發生了車禍。雖然大難不死，可是眼睛完全失明，什麼也看不見了。他進入一家盲人重建院，在那裡學習如何用點字技巧，如何靠手杖走路等等。慢慢地終於能夠獨立生活了。

他說：「可是我最苦惱的是，我弄不清楚對方是不是黑人。」我向我的心理輔導員談我的問題，他也儘量開導我，我非常信賴他，什麼都告訴他，將他看成自己的良師益友。

有一天，那位輔導員告訴我，他本人就是位黑人。從此以後，我的偏見就慢慢完全消失了，我看不出人是白人，還是黑人。對我來講，我只知道他是好人，還是壞人；至於膚色，對我已絕對地無意義了。

車子快到波士頓，老先生說：「我失去了視力，也失去了偏見，多麼幸福的事！」

在月台上，老先生的太太已在等他，兩人親切地擁抱。我赫然發現他太太是一位滿頭銀髮的黑人，當時大吃一驚。 

我這才發現，我視力良好，因此我偏見猶在，多麼不幸的事！"""

questions = [
    {
        "id": 1,
        "type": "句子挖空選擇題",
        "question": "「當時正值洛杉磯種族暴動的時期，我們的談話______就談到了種族偏見的問題。」",
        "options": ["(A) 因而", "(B) 於是", "(C) 因此", "(D) 可是"],
        "answer": "C",
        "explanation": "根據課文原文：「我們的談話因此就談到了種族偏見的問題。」此處使用「因此」表示因果關係。"
    },
    {
        "id": 2,
        "type": "句子挖空選擇題",
        "question": "「在南方這句話就是『我們不歡迎黑人』的意思，當時舉班______，他還被系主任抓去罵了一頓。」",
        "options": ["(A) 歡騰", "(B) 譁然", "(C) 雀躍", "(D) 寂靜"],
        "answer": "B",
        "explanation": "原文為「當時舉班譁然」，「譁然」形容眾人議論紛紛、喧鬧的樣子。"
    },
    {
        "id": 3,
        "type": "句子挖空選擇題",
        "question": "「我不和他們來往，如何會和黑人結婚？說實話，我當時認為任何白人和黑人結婚都會使父母______。」",
        "options": ["(A) 蒙辱", "(B) 蒙恩", "(C) 光彩", "(D) 驕傲"],
        "answer": "A",
        "explanation": "根據原文，盲人老先生年輕時有嚴重的種族偏見，認為與黑人結婚會使父母「蒙辱」（蒙受恥辱）。"
    },
    {
        "id": 4,
        "type": "句子挖空選擇題",
        "question": "「雖然大難不死，可是眼睛完全______，什麼也看不見了。」",
        "options": ["(A) 昏暗", "(B) 模糊", "(C) 失明", "(D) 疲憊"],
        "answer": "C",
        "explanation": "車禍後，老先生眼睛完全「失明」，代表喪失視力。"
    },
    {
        "id": 5,
        "type": "句子挖空選擇題",
        "question": "「我向我的心理輔導員談我的問題，他也儘量開導我，我非常信賴他，將他看成自己的______。」",
        "options": ["(A) 競爭對手", "(B) 莫逆之交", "(C) 良師益友", "(D) 萍水相逢"],
        "answer": "C",
        "explanation": "原文為「將他看成自己的良師益友」，指能給予良好教導與益處的朋友。"
    },
    {
        "id": 6,
        "type": "句子挖空選擇題",
        "question": "「在月台上，老先生的太太已在等他，兩人親切地擁抱。我______發現他太太是一位滿頭銀髮的黑人...」",
        "options": ["(A) 赫然", "(B) 突然", "(C) 依然", "(D) 果然"],
        "answer": "A",
        "explanation": "根據原文，「我赫然發現他太太是一位滿頭銀髮的黑人」，「赫然」表示令人吃驚、驚奇的樣子。"
    },
    {
        "id": 7,
        "type": "文意理解選擇題",
        "question": "根據本文，作者為什麼和盲人談起話來一點困難也沒有？",
        "options": ["(A) 作者自己也是盲人", "(B) 作者的博士論文指導教授是位盲人", "(C) 作者曾做過盲人重建院的義工", "(D) 作者擅長點字溝通"],
        "answer": "B",
        "explanation": "課文開頭提到：「我的博士論文指導教授是位盲人，因此我和盲人談起話來，一點困難也沒有。」"
    },
    {
        "id": 8,
        "type": "文意理解選擇題",
        "question": "老先生在南方長大時，對黑人抱持著什麼樣的態度？",
        "options": ["(A) 認為黑人與白人完全平等", "(B) 認為黑人應享有更多特權", "(C) 認為黑人低人一等，並刻意避免接觸", "(D) 積極幫助黑人爭取權益"],
        "answer": "C",
        "explanation": "老先生自述從小認為黑人低人一等，不與黑人同桌吃飯、不上同所學校，甚至付錢時將錢放櫃台上以避免手部接觸。"
    },
    {
        "id": 9,
        "type": "文意理解選擇題",
        "question": "老先生在請帖上寫著「我們保留拒絕任何人的權利」，這句話在當時南方的真實含意是什麼？",
        "options": ["(A) 我們不歡迎黑人", "(B) 我們歡迎所有人參加", "(C) 必須憑票入場", "(D) 必須穿著正式服裝"],
        "answer": "A",
        "explanation": "課文提到：「在南方這句話就是『我們不歡迎黑人』的意思」。"
    },
    {
        "id": 10,
        "type": "文意理解選擇題",
        "question": "老先生是如何開始消除他對黑人的偏見的？",
        "options": ["(A) 讀了許多關於平權運動的書籍", "(B) 車禍後在盲人重建院被一位黑人輔導員開導，且他極其信賴對方", "(C) 太太說服他改變看法", "(D) 他在暴動中被黑人所救"],
        "answer": "B",
        "explanation": "老先生極起信賴與感激的心理輔導員告訴他自己是黑人，這讓老先生的偏見慢慢完全消失。"
    },
    {
        "id": 11,
        "type": "文意理解選擇題",
        "question": "老先生說：「我失去了視力，也失去了偏見，多麼幸福的事！」這句話代表他此時的心境如何？",
        "options": ["(A) 雖然看不見，但心靈從偏見的枷鎖中解放，感到無比輕鬆與幸福", "(B) 因為失明而自暴自棄", "(C) 抱怨命運捉弄，讓他看不見美麗的世界", "(D) 慶幸自己不用再看到黑人"],
        "answer": "A",
        "explanation": "失去視力使他無法再以膚色判斷人，進而消除了偏見，使他能以純粹的心結識好人，因此他感到幸福。"
    },
    {
        "id": 12,
        "type": "文意理解選擇題",
        "question": "作者在火車旅程結束時，看見老先生的太太是位黑人，當時的心情是什麼？",
        "options": ["(A) 憤怒與不解", "(B) 欣慰與高興", "(C) 大吃一驚，並反思自己的偏見", "(D) 漠不關心"],
        "answer": "C",
        "explanation": "作者看到老先生太太是黑人時大吃一驚，並得出結論：「我這才發現，我視力良好，因此我偏見猶在，多麼不幸的事！」"
    },
    {
        "id": 13,
        "type": "文意理解選擇題",
        "question": "關於本文的寫作手法，下列敘述何者正確？",
        "options": ["(A) 透過大量抽象的學術理論來探討種族偏見", "(B) 藉由與盲人老先生的巧遇和談話，以故事性的敘述帶出深刻的省思", "(C) 以第三人稱全知觀點，詳細描繪洛杉磯暴動的過程", "(D) 採用倒敘法，先寫老先生的太太，再寫火車上的談話"],
        "answer": "B",
        "explanation": "本文採用第一人稱（「我」），透過在火車上與盲人老先生的對話，以及最後月台上的驚訝發現，以故事形式引導讀者反思「視力與偏見」的關係。"
    },
    {
        "id": 14,
        "type": "素養選擇題",
        "question": "本文標題為「視力與偏見」，其中「視力」在文章最後的象徵意義最接近下列何者？",
        "options": ["(A) 探索世界的工具", "(B) 阻礙真實理解、造成刻板印象的外在表象", "(C) 盲人老先生最大的遺憾", "(D) 指導教授的博學多聞"],
        "answer": "B",
        "explanation": "作者最後提到「我視力良好，因此我偏見猶在」，此處的「視力」象徵容易讓人只看外表（如膚色、外貌）而產生主觀偏見的感官限制。"
    },
    {
        "id": 15,
        "type": "素養選擇題",
        "question": "盲人老先生說：「我看不出人是白人，還是黑人。對我來講，我只知道他是好人，還是壞人；至於膚色，對我已絕對地無意義了。」這句話傳達的核心價值是什麼？",
        "options": ["(A) 應以一個人的內在品德而非外在特徵來評斷人", "(B) 世界上只有絕對的好人與壞人", "(C) 人類不應該發展視覺科技", "(D) 盲人比明眼人更適合當心理輔導員"],
        "answer": "A",
        "explanation": "老先生不再受膚色干擾，只專注於人的內在好壞，強調評價他人時應注重內在本質而非外在標籤。"
    },
    {
        "id": 16,
        "type": "素養選擇題",
        "question": "下列哪一個社會現象，最符合本文所探討的「偏見猶在」的盲點？",
        "options": ["(A) 為了環保，大家出門自備購物袋和環保杯", "(B) 企業面試時，只因為求職者的性別或年齡就直接將履歷淘汰，不看其專業能力", "(C) 政府積極建設盲人步道與無障礙空間", "(D) 學校舉辦多元文化週，介紹各國傳統美食"],
        "answer": "B",
        "explanation": "選項反映了以刻板印象（年齡、性別）代替對個人真實能力（內在）的評估，這正是「偏見猶在」的體現。"
    },
    {
        "id": 17,
        "type": "素養選擇題",
        "question": "故事中，老先生在失明後，直到心理輔導員向他表明自己是黑人，他才徹底消除偏見。這說明了偏見的消除往往需要什麼關鍵？",
        "options": ["(A) 強迫性的法律限制", "(B) 長期的自我封閉", "(C) 建立在真誠信任與深刻理解基礎上的生命經驗", "(D) 遭遇重大車禍的生理打擊"],
        "answer": "C",
        "explanation": "老先生非常信賴該輔導員，視其為良師益友，當得知他是黑人後，其根深蒂固的種族偏見與對該黑人朋友的真實喜愛發生碰撞，最終偏見才得以消除。這說明了真實、真誠的交往經驗對消除偏見的重要性。"
    },
    {
        "id": 18,
        "type": "素養選擇題",
        "question": "英國作家王爾德曾說：「只有膚淺的人，才不以貌取人。」若對照本文，本文作者的觀點與此話的關係為何？",
        "options": ["(A) 兩者觀點完全一致，都鼓勵人們多看外表", "(B) 本文作者持相反觀點，認為「以貌取人」（如看膚色）正是偏見的來源", "(C) 王爾德的「以貌取人」是指品味，與本文無關，無法對比", "(D) 本文作者認為「以貌取人」是盲人的專利"],
        "answer": "B",
        "explanation": "本文強調明眼人常因看到外在（膚色）而產生偏見，因此倡導應超越外在，看重內在本質，這與字面上的「以貌取人」相反。"
    },
    {
        "id": 19,
        "type": "素養選擇題",
        "question": "如果我們要為這篇課文設計一個討論課的主題，下列哪一個最切合文章旨意？",
        "options": ["(A) 「如何預防鐵路行車事故」", "(B) 「眼睛看見的，就是真實嗎？——談偏見的形成與超越」", "(C) 「美國南方的歷史與地理變遷」", "(D) 「盲人重建院的點字與復健技術」"],
        "answer": "B",
        "explanation": "本文探討「視力好反而容易產生偏見，失去視力反而消除偏見」，主題「眼睛看見的，就是真實嗎？——談偏見的形成與超越」完美扣合這一主旨。"
    },
    {
        "id": 20,
        "type": "素養選擇題",
        "question": "課文最後，作者說「我這才發現，我視力良好，因此我偏見猶在，多麼不幸的事！」這句話的語氣充滿了什麼？",
        "options": ["(A) 自豪與驕傲", "(B) 反諷與深沉的自我省思", "(C) 嫉妒與不甘", "(D) 悲觀與絕望"],
        "answer": "B",
        "explanation": "作者用「多麼不幸的事」來形容自己擁有健康的視力，這是一種深刻的反思與警惕，表示自己仍受外在表象干擾而留有偏見，具有強烈的省思意味。"
    }
]

questions_json = json.dumps(questions, ensure_ascii=False)

# Group definition
groups = {
    "一、句子挖空選擇題 (共 6 題，每題 5 分)": [q for q in questions if q["type"] == "句子挖空選擇題"],
    "二、文意理解選擇題 (共 7 題，每題 5 分)": [q for q in questions if q["type"] == "文意理解選擇題"],
    "三、素養選擇題 (共 7 題，每題 5 分)": [q for q in questions if q["type"] == "素養選擇題"]
}

# Helper function to split text into .read-sentence spans
def split_to_spans(text):
    parts = re.split(r'([_＿]+)', text)
    html_parts = []
    for part in parts:
        if re.match(r'^[_＿]+$', part):
            html_parts.append(part)
        else:
            if part.strip():
                subparts = re.split(r'([，。？！；：\s\n]+)', part)
                buffer = ""
                for sub in subparts:
                    if re.match(r'^[，。？！；：\s\n]+$', sub):
                        if buffer:
                            clean_buf = buffer.strip().replace('"', '&quot;').replace("'", '&#39;')
                            if clean_buf:
                                html_parts.append(f'<span class="read-sentence" onclick="speakSentence(this, event)">{clean_buf}</span>')
                            buffer = ""
                        html_parts.append(sub)
                    else:
                        if buffer:
                            clean_buf = buffer.strip().replace('"', '&quot;').replace("'", '&#39;')
                            if clean_buf:
                                html_parts.append(f'<span class="read-sentence" onclick="speakSentence(this, event)">{clean_buf}</span>')
                            buffer = ""
                        buffer = sub
                if buffer:
                    clean_buf = buffer.strip().replace('"', '&quot;').replace("'", '&#39;')
                    if clean_buf:
                        html_parts.append(f'<span class="read-sentence" onclick="speakSentence(this, event)">{clean_buf}</span>')
            else:
                html_parts.append(part)
    return "".join(html_parts)

# Generate student questions list HTML (grouped by section without question type prefixes)
student_questions_list_html = []
global_idx = 1
for group_name, group_qs in groups.items():
    student_questions_list_html.append(f"<h3 style='margin-top: 25px; border-bottom: 2px solid var(--primary-color); padding-bottom: 5px;'>{group_name}</h3>")
    for q in group_qs:
        q_spans = split_to_spans(q["question"])
        item_html = f"""
        <div class="question-card" id="q-{q['id']}">
          <div class="question-title">
            <button class="speaker-btn" onclick="speakQuestion(this, 'q-{q['id']}')">🔊</button>
            <span>{global_idx}. </span>
            <span style="display:inline-block; margin-left:5px;">{q_spans}</span>
          </div>
          <ul class="options-list">
        """
        for opt in q["options"]:
            opt_spans = split_to_spans(opt)
            input_id = f"q-{q['id']}-opt-{opt[1]}"
            item_html += f"""
            <li class="option-item">
              <input type="radio" name="q-{q['id']}" id="{input_id}">
              <label for="{input_id}">{opt_spans}</label>
            </li>
            """
        item_html += """
          </ul>
        </div>
        """
        student_questions_list_html.append(item_html)
        global_idx += 1
student_questions_html = "\n".join(student_questions_list_html)


# Read Student HTML Template or fallback
student_html_template = ""
if os.path.exists(student_template_file):
    print("Loading student HTML template from configuration...")
    with open(student_template_file, "r", encoding="utf-8") as f:
        student_html_template = f.read()
else:
    print("Student template not found. Using default fallback...")
    # Default fallback string (in case the file is missing)
    student_html_template = """<!DOCTYPE html>
<html lang="zh-TW">
<head>
  <meta charset="UTF-8">
  <title>{{EXAM_TITLE}}</title>
  <link href="https://fonts.googleapis.com/css2?family=LXGW+WenKai+TC&display=swap" rel="stylesheet">
  <style>
    :root {
      --bg-color: #fafaf9;
      --card-bg: #ffffff;
      --text-color: #1c1917;
      --primary-color: #3b82f6;
      --border-color: #e7e5e4;
      --toolbar-bg: rgba(255, 255, 255, 0.85);
      --highlight-bg: #fef08a;
      --font-size: 14pt;
    }
    body {
      background-color: var(--bg-color);
      color: var(--text-color);
      font-family: 'LXGW WenKai TC', sans-serif;
      margin: 0;
      padding: 0;
      font-size: var(--font-size);
    }
    .container { max-width: 900px; margin: 30px auto; padding: 0 20px; }
    .toolbar { position: sticky; top: 0; background: var(--toolbar-bg); padding: 10px; display: flex; justify-content: space-between; }
    .read-sentence { cursor: pointer; border-bottom: 1px dashed transparent; }
    .read-sentence:hover { border-bottom-color: var(--primary-color); }
    .speaking-highlight { background-color: var(--highlight-bg) !important; color: #000 !important; }
  </style>
</head>
<body>
  <div class="toolbar">
    <button onclick="window.print()">列印</button>
  </div>
  <div class="container">
    <h1>{{EXAM_TITLE}}</h1>
    <div id="questions-container">
      {{QUESTIONS_HTML}}
    </div>
  </div>
</body>
</html>"""

# Do replacements
student_html_content_final = (student_html_template
                              .replace("{{EXAM_TITLE}}", "國文科評量《視力與偏見》學生版練習卷")
                              .replace("{{TOPIC}}", article_title)
                              .replace("{{QUESTIONS_COUNT}}", str(len(questions)))
                              .replace("{{QUESTIONS_HTML}}", student_questions_html))

with open(student_html_path, "w", encoding="utf-8") as f:
    f.write(student_html_content_final)
print(f"Created student html: {student_html_path}")


# Read Teacher Slideshow Template or fallback
teacher_html_template = ""
if os.path.exists(teacher_template_file):
    print("Loading teacher HTML template from configuration...")
    with open(teacher_template_file, "r", encoding="utf-8") as f:
        teacher_html_template = f.read()
else:
    print("Teacher template not found. Using default fallback...")
    teacher_html_template = """<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <title>教師簡報</title>
</head>
<body>
  <div id="question-text"></div>
  <script>
    const questions = {{QUESTIONS_JSON}};
    console.log(questions);
  </script>
</body>
</html>"""

teacher_html_content_final = (teacher_html_template
                              .replace("{{TOPIC}}", article_title)
                              .replace("{{QUESTIONS_JSON}}", questions_json))

with open(teacher_html_path, "w", encoding="utf-8") as f:
    f.write(teacher_html_content_final)
print(f"Created teacher slideshow html: {teacher_html_path}")


# PDF GENERATION LOGIC
def make_pdf(is_teacher):
    suffix = "解析" if is_teacher else "練習卷"
    temp_html_path = os.path.join(TARGET_DIR, f"temp_print_{suffix}.html")
    pdf_output_path = os.path.join(TARGET_DIR, f"視力與偏見_20題_{suffix}.pdf")
    
    questions_list_html = []
    global_idx = 1
    for group_name, group_qs in groups.items():
        questions_list_html.append(f"<h3 style='margin-top: 25px; border-bottom: 2px solid #000; padding-bottom: 5px;'>{group_name}</h3>")
        for q in group_qs:
            q_html = f"""
            <div style="margin-bottom: 20px; page-break-inside: avoid;">
              <strong>{global_idx}. {q['question']}</strong>
              <div style="margin-left: 20px; margin-top: 5px;">
            """
            for opt in q["options"]:
                q_html += f"<div style='margin-bottom: 4px;'>{opt}</div>"
            
            q_html += "</div>"
            
            if is_teacher:
                q_html += f"""
                <div style="margin-top: 8px; color: #b91c1c; font-weight: bold; margin-left: 20px;">
                  【正確答案】：{q['answer']}
                </div>
                <div style="color: #4b5563; font-style: italic; margin-left: 20px; margin-top: 3px;">
                  【詳細解析】：{q['explanation']}
                </div>
                """
            q_html += "</div>"
            questions_list_html.append(q_html)
            global_idx += 1
    
    questions_content = "\n".join(questions_list_html)
    
    temp_html_content = f"""<!DOCTYPE html>
    <html lang="zh-TW">
    <head>
      <meta charset="UTF-8">
      <title>視力與偏見_20題_{suffix}</title>
      <link href="https://fonts.googleapis.com/css2?family=LXGW+WenKai+TC&display=swap" rel="stylesheet">
      <style>
        body {{
          font-family: 'LXGW WenKai TC', 'Microsoft JhengHei', sans-serif;
          font-size: 14pt; /* 14pt requested */
          line-height: 1.5;
          color: #000000;
          background: #ffffff;
        }}
        @page {{
          size: A4;
          margin: 1.8cm 1.5cm;
        }}
        @media print {{
          body {{
            margin: 0;
          }}
        }}
      </style>
    </head>
    <body>
      <div style="text-align: center; margin-bottom: 20px;">
        <h2 style="margin: 0;">國文科課文評量《視力與偏見》{suffix}</h2>
        <p style="margin: 5px 0;">班級：___________ 座號：______ 姓名：___________ 得分：___________</p>
        <hr style="border: 0; border-top: 2px solid #000;">
      </div>
      
      <div style="margin-top: 20px;">
        {questions_content}
      </div>
    </body>
    </html>
    """
    
    with open(temp_html_path, "w", encoding="utf-8") as f:
        f.write(temp_html_content)
        
    edge_path = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
    cmd = [
        edge_path,
        "--headless",
        "--disable-gpu",
        "--no-sandbox",
        "--no-pdf-header-footer",
        f"--print-to-pdf={pdf_output_path}",
        temp_html_path
    ]
    print(f"Generating PDF for: {suffix}...")
    subprocess.run(cmd, capture_output=True)
    
    # Delete temporary HTML file
    if os.path.exists(temp_html_path):
        os.remove(temp_html_path)

make_pdf(is_teacher=False)
make_pdf(is_teacher=True)


# DOCX GENERATION LOGIC
def make_docx(is_teacher):
    suffix = "解析" if is_teacher else "練習卷"
    docx_output_path = os.path.join(TARGET_DIR, f"視力與偏見_20題_{suffix}.docx")
    
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    def format_run(run, name="微軟正黑體", size_pt=14, bold=False, italic=False, color_rgb=None):
        run.font.name = name
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), name)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rPr.append(rFonts)
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"國文科課文評量《視力與偏見》{suffix}")
    format_run(run_title, name="微軟正黑體", size_pt=18, bold=True)
    
    # Header Info Line
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = p_info.add_run("班級：___________  座號：______  姓名：___________  得分：___________")
    format_run(run_info, name="微軟正黑體", size_pt=12)
    
    # Separator
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = p_sep.add_run("―" * 40)
    format_run(run_sep, size_pt=10)

    # Questions and Sections
    global_idx = 1
    for group_name, group_qs in groups.items():
        p_group = doc.add_paragraph()
        p_group.paragraph_format.space_before = Pt(18)
        p_group.paragraph_format.space_after = Pt(6)
        run_group = p_group.add_run(group_name)
        format_run(run_group, name="微軟正黑體", size_pt=14, bold=True)
        
        for q in group_qs:
            p_q = doc.add_paragraph()
            p_q.paragraph_format.space_before = Pt(8)
            run_q = p_q.add_run(f"{global_idx}. {q['question']}")
            format_run(run_q, name="微軟正黑體", size_pt=14, bold=True)
            
            for opt in q["options"]:
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.left_indent = Inches(0.4)
                run_opt = p_opt.add_run(opt)
                format_run(run_opt, name="微軟正黑體", size_pt=14)
                
            if is_teacher:
                p_ans = doc.add_paragraph()
                p_ans.paragraph_format.left_indent = Inches(0.4)
                p_ans.paragraph_format.space_before = Pt(4)
                run_ans = p_ans.add_run(f"【正確答案】：{q['answer']}")
                format_run(run_ans, name="微軟正黑體", size_pt=14, bold=True, color_rgb=RGBColor(185, 28, 28))
                
                p_exp = doc.add_paragraph()
                p_exp.paragraph_format.left_indent = Inches(0.4)
                p_exp.paragraph_format.space_after = Pt(8)
                run_exp = p_exp.add_run(f"【詳細解析】：{q['explanation']}")
                format_run(run_exp, name="微軟正黑體", size_pt=12, italic=True, color_rgb=RGBColor(75, 85, 99))
            
            global_idx += 1
            
    doc.save(docx_output_path)
    print(f"Created docx: {docx_output_path}")

make_docx(is_teacher=False)
make_docx(is_teacher=True)

print("All 6 educational media files successfully generated.")
